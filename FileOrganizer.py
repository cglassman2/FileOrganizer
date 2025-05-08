import sys
import os
import subprocess
from PyQt6.QtWidgets import (
QApplication, QWidget, QVBoxLayout, QHBoxLayout, QPushButton, QFileDialog,
QTreeWidget, QTreeWidgetItem, QLabel, QGridLayout, QScrollArea, QFrame,
QListWidget, QMessageBox, QTabWidget, QSplitter, QGroupBox, QComboBox,
QInputDialog, QTableWidget, QTableWidgetItem, QHeaderView, QCheckBox, QLineEdit,
QTextBrowser,  QAbstractItemView, QAbstractScrollArea, QMenu)
from PyQt6.QtGui import QFileSystemModel, QDrag, QAction, QClipboard
import time
from PyQt6.QtCore import Qt, QThread, pyqtSignal, QTimer,QMimeData, QUrl, QByteArray, QTimer
from pathlib import Path
import shutil
from docx import Document
from datetime import datetime
import re
from urllib.parse import urlparse


class FileLoaderThread(QThread):
    files_loaded = pyqtSignal(list)

    def __init__(self, path, parent=None):
        super().__init__(parent)
        self.path = path

    def run(self):
        files = []
        try:
            with os.scandir(self.path) as entries:
                for entry in entries:
                    if entry.is_file():
                        files.append(entry.path)
                        if len(files) >= 700:  # Only load first 500 paths
                            break
        except Exception as e:
            print(f"[ERROR] Failed to load files in background: {e}")
        self.files_loaded.emit(files)


class BreadcrumbLabel(QTextBrowser):
    """Custom QTextBrowser that shows clickable breadcrumb links."""
    def __init__(self, file_explorer_app, parent=None):
        super().__init__(parent)
        self.file_explorer_app = file_explorer_app
        self.setOpenExternalLinks(False)
        self.setOpenLinks(False)
        self.setStyleSheet("border: none; background: transparent;")
        self.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setVerticalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.setMaximumHeight(15)
        self.anchorClicked.connect(lambda url: self.file_explorer_app.on_breadcrumb_clicked(url.toString()))

    def setBreadcrumbs(self, path_list):
        """Render breadcrumb trail from a list of (name, absolute_path) pairs."""
        from urllib.parse import quote

        breadcrumb_links = []
        for name, path in path_list:
            abs_path = os.path.abspath(path)
            href = "file:///" + quote(abs_path.replace(os.sep, "/"))
            breadcrumb_links.append(f'<a href="{href}">{name}</a>')

        html = ' / '.join(breadcrumb_links)
        self.setHtml(html)



class FileExplorerApp(QWidget):
    BOOKMARKS_FILE = "bookmarks.txt"
    NOTES_FILE = "notes.txt"
    SAVED_FILES_FILE = "saved_files_all.txt"
    if getattr(sys, 'frozen', False):
        BASE_DIR = os.path.dirname(sys.executable)
    else:
        BASE_DIR = os.path.dirname(os.path.abspath(__file__))

    LISTS_DIR = os.path.join(BASE_DIR, "lists")



    def __init__(self):
        super().__init__()
        self.setWindowTitle("File Organizer")
        self.setGeometry(100, 100, 1550, 850)
        
        self.selected_files = []
        self.notes = {}


        # Create tab widget
        self.tabs = QTabWidget()
        self.layout = QVBoxLayout()
        self.layout.addWidget(self.tabs)
        self.setLayout(self.layout)


        # First tab: File Explorer
        self.tab1 = QWidget()
        self.setup_file_explorer_tab()
        self.tabs.addTab(self.tab1, "File Explorer")

        # Add widgets to splitter
        self.splitter.addWidget(self.bookmark_section)  # Bookmarks section
        self.splitter.addWidget(self.tree_widget)       # File tree section
        self.splitter.addWidget(self.scroll_area)       # File grid section

        self.files_widget.setAcceptDrops(True)
        self.files_widget.dragEnterEvent = self.dragEnterEvent
        self.files_widget.dragMoveEvent = self.dragMoveEvent
        self.files_widget.dropEvent = self.dropEvent
       


        # Set default heights (e.g., 150 px for bookmarks, 300 px for tree, 400 px for file grid)
        self.splitter.setSizes([150, 125, 400])

        # Second tab: Saved Files
        self.tab2 = QWidget()
        self.setup_saved_files_tab()
        self.tabs.addTab(self.tab2, "Saved Files")

        # QFileSystemModel for file icons
        self.file_model = QFileSystemModel()
        self.file_model.setRootPath("")

        # Current directory and file storage
        self.current_directory = None


        # Set default directory to user's Documents
        self.current_directory = str(Path.home() / "Documents")

        # Load that directory into the UI
        self.update_breadcrumb(self.current_directory)
        self.populate_tree(self.current_directory)

        # Safely list files
        try:
            self.all_files = [os.path.join(self.current_directory, f) for f in os.listdir(self.current_directory)]
        except Exception as e:
            print(f"[ERROR] Could not list files in default directory: {e}")
        self.display_files()
        #self.all_files = []

        #self.load_bookmarks()
        self.load_notes()  # ✅ load notes early
        self.load_sections(load_notes=False)  # don't reload notes again here
        self.update_files_table()  # now safe



    
    def setup_file_explorer_tab(self):
        layout = QVBoxLayout()
        layout.setContentsMargins(5, 5, 5, 5)
        layout.setSpacing(5)

        # --- Single Horizontal Control Bar ---
        control_row = QHBoxLayout()
        control_row.setSpacing(10)

        self.change_directory_button = QPushButton("Change Directory")
        self.change_directory_button.clicked.connect(self.open_directory_dialog)
        control_row.addWidget(self.change_directory_button)


        sort_label = QLabel("Sort Bookmarks:")
        sort_label.setFixedWidth(90)
        control_row.addWidget(sort_label)

        self.sort_bookmark_combo = QComboBox()
        self.sort_bookmark_combo.addItems([
            "Sort by Name (A-Z)",
            "Sort by Name (Z-A)",
            "Sort by Depth (Shallow to Deep)",
            "Sort by Depth (Deep to Shallow)",
            "Sort by Last Modified (Newest First)",
            "Sort by Last Modified (Oldest First)",
            "Sort by Path Length (Shortest First)",
            "Sort by Path Length (Longest First)"
        ])
        self.sort_bookmark_combo.currentIndexChanged.connect(self.sort_bookmarks)
        self.sort_bookmark_combo.setFixedSize(270, 30)
        control_row.addWidget(self.sort_bookmark_combo)


        self.bookmark_button = QPushButton("Bookmark Current Folder")
        self.bookmark_button.clicked.connect(self.add_bookmark)
        control_row.addWidget(self.bookmark_button)

        self.remove_bookmark_button = QPushButton("Remove Selected Bookmark")
        self.remove_bookmark_button.clicked.connect(self.remove_selected_bookmark)
        control_row.addWidget(self.remove_bookmark_button)
        
        self.save_selected_button = QPushButton("Save Selected File(s)")
        self.save_selected_button.clicked.connect(self.save_selected_files)
        control_row.addWidget(self.save_selected_button)

        section_label = QLabel("Save File(s) To:")
        section_label.setFixedWidth(72)
        control_row.addWidget(section_label)

        self.section_combo_file_explorer = QComboBox()
        self.section_combo_file_explorer.setEditable(True)
        self.section_combo_file_explorer.setFixedSize(150, 30)
        self.section_combo_file_explorer.setStyleSheet("""
            QComboBox {
                font-size: 13px;
                font-weight: bold;
            }
        """)
        control_row.addWidget(self.section_combo_file_explorer)
        

        layout.addLayout(control_row)

        # --- Splitter: Bookmarks | Tree | Files ---
        self.splitter = QSplitter(Qt.Orientation.Horizontal)

        # Bookmark Section
        self.bookmark_section = QGroupBox("Bookmarks")
        bookmark_layout = QVBoxLayout()

        self.bookmark_list = QTableWidget()
        self.bookmark_list.setColumnCount(1)
        self.bookmark_list.setHorizontalHeaderLabels(["Bookmarks"])
        self.bookmark_list.verticalHeader().setVisible(False)
        self.bookmark_list.horizontalHeader().setVisible(False)
        self.bookmark_list.setShowGrid(False)
        self.bookmark_list.setEditTriggers(QTableWidget.EditTrigger.NoEditTriggers)
        self.bookmark_list.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectItems)
        self.bookmark_list.cellClicked.connect(self.open_bookmarked_folder_from_table)
        self.bookmark_list.itemClicked.connect(self.open_bookmarked_folder)
        self.bookmark_list.setColumnWidth(0, 225)
        self.bookmark_list.horizontalHeader().setSectionResizeMode(QHeaderView.ResizeMode.Interactive)

        bookmark_layout.addWidget(self.bookmark_list)
        self.bookmark_section.setLayout(bookmark_layout)
        self.splitter.addWidget(self.bookmark_section)
        self.splitter.setCollapsible(0, True)
        
        # Breadcrumb path
        self.path_label = BreadcrumbLabel(self)
        self.path_label.setTextInteractionFlags(Qt.TextInteractionFlag.TextBrowserInteraction)
        self.path_label.setOpenExternalLinks(True)
        layout.addWidget(self.path_label)

        # Tree View
        self.tree_widget = QTreeWidget()
        self.tree_widget.setHeaderLabels(["Files & Folders"])
        self.tree_widget.itemClicked.connect(self.on_item_clicked)
        self.tree_widget.itemDoubleClicked.connect(self.on_item_double_clicked)
        self.tree_widget.itemExpanded.connect(self.on_tree_item_expanded)
        self.splitter.addWidget(self.tree_widget)

        # File grid
        self.scroll_area = QScrollArea()
        self.scroll_area.setWidgetResizable(True)
        self.files_widget = QWidget()
        self.grid_layout = QGridLayout(self.files_widget)
        self.scroll_area.setWidget(self.files_widget)
        self.splitter.addWidget(self.scroll_area)

        layout.addWidget(self.splitter)

        # Bottom view options row
        view_options_layout = QHBoxLayout()
        

        self.copy_files_checkbox = QCheckBox("When dragging files, check to copy, uncheck to move.")
        self.copy_files_checkbox.setChecked(True)
        view_options_layout.addWidget(self.copy_files_checkbox)

        self.search_box_explorer = QLineEdit()
        self.search_box_explorer.setPlaceholderText("Search files...")
        self.search_box_explorer.textChanged.connect(self.filter_explorer_files)
        view_options_layout.addWidget(self.search_box_explorer)
        
        self.view_mode_combo = QComboBox()
        self.view_mode_combo.addItems(["Detailed View", "Icon View", "List View"])
        self.view_mode_combo.currentIndexChanged.connect(self.display_files)
        view_options_layout.addWidget(self.view_mode_combo)

        self.sort_combo_explorer = QComboBox()
        self.sort_combo_explorer.addItems([
            "Sort by Date Created", "Sort by Date Modified", "Sort by Date Accessed"
        ])
        self.sort_combo_explorer.currentIndexChanged.connect(self.sort_explorer_files)
        view_options_layout.addWidget(self.sort_combo_explorer)

        layout.addLayout(view_options_layout)

        self.tab1.setLayout(layout)

        # Load bookmarks
        self.load_bookmarks()


    def on_item_double_clicked(self, item, column):
        path = item.data(0, Qt.ItemDataRole.UserRole)
        if os.path.isdir(path):
            # Update internal state
            self.current_directory = path
            self.update_breadcrumb(path)
            self.load_thread = FileLoaderThread(path)
            self.load_thread.files_loaded.connect(self.on_files_loaded)
            self.load_thread.start()

            # And open in system file explorer
            self.open_folder_in_explorer(path)

    def sort_bookmarks(self):
        all_bookmarks = []

        # Gather all bookmarks from the table
        for row in range(self.bookmark_list.rowCount()):
            for col in range(self.bookmark_list.columnCount()):
                item = self.bookmark_list.item(row, col)
                if item:
                    all_bookmarks.append(item.text())

        sort_type = self.sort_bookmark_combo.currentText()

        def get_depth(path):
            return path.count(os.sep)

        def get_mtime(path):
            try:
                return os.path.getmtime(path)
            except Exception:
                return float('-inf')  # Push missing paths to end

        if "Name (A-Z)" in sort_type:
            all_bookmarks.sort()
        elif "Name (Z-A)" in sort_type:
            all_bookmarks.sort(reverse=True)
        elif "Depth (Shallow to Deep)" in sort_type:
            all_bookmarks.sort(key=get_depth)
        elif "Depth (Deep to Shallow)" in sort_type:
            all_bookmarks.sort(key=get_depth, reverse=True)
        elif "Modified (Newest First)" in sort_type:
            all_bookmarks.sort(key=get_mtime, reverse=True)
        elif "Modified (Oldest First)" in sort_type:
            all_bookmarks.sort(key=get_mtime)
        elif "Path Length (Shortest First)" in sort_type:
            all_bookmarks.sort(key=lambda x: len(x))
        elif "Path Length (Longest First)" in sort_type:
            all_bookmarks.sort(key=lambda x: len(x), reverse=True)
            
        self.save_bookmarks_from_list(all_bookmarks)
        self.load_bookmarks()


    def open_bookmarked_folder_from_table(self, row, col):
        item = self.bookmark_list.item(row, col)
        if item:
            path = item.text()
            if os.path.exists(path):
                self.current_directory = path
                self.update_breadcrumb(path)
                self.populate_tree(path)
                self.load_thread = FileLoaderThread(path)
                self.load_thread.files_loaded.connect(self.on_files_loaded)
                self.load_thread.start()

    def on_tree_item_expanded(self, item):
        if item.childCount() == 1 and item.child(0).text(0) == "Loading...":
            item.removeChild(item.child(0))  # Remove dummy
            path = item.data(0, Qt.ItemDataRole.UserRole)
            self.populate_subitems(item, path)


    def filter_explorer_files(self):
        """Filter files based on search query for all views."""
        query = self.search_box_explorer.text().strip().lower()
        view_mode = self.view_mode_combo.currentText()

        if view_mode == "Detailed View":
            # Get the table widget from the grid
            table = self.grid_layout.itemAt(0).widget()
            if isinstance(table, QTableWidget):
                for row in range(table.rowCount()):
                    file_name = table.item(row, 0).text().lower()
                    table.setRowHidden(row, query not in file_name)
        else:
            # Handle icon/list view filtering
            for i in range(self.grid_layout.count()):
                widget = self.grid_layout.itemAt(i).widget()
                if widget:
                    file_labels = widget.findChildren(QLabel)
                    if len(file_labels) >= 2:
                        file_label = file_labels[1]
                        file_name = file_label.text()
                        if "<a href=" in file_name:
                            file_name = file_name.split(">")[1].split("</a")[0]
                        file_name = file_name.strip().lower()
                        widget.setVisible(query in file_name)



    def setup_saved_files_tab(self):
        layout = QVBoxLayout()
        button_layout = QHBoxLayout()

        self.section_combo = QComboBox()
        self.section_combo.setEditable(True)
        self.section_combo.lineEdit().setPlaceholderText("")
        self.section_combo.currentTextChanged.connect(self.load_files_for_section)
        self.section_combo.setFixedSize(150, 30)
        self.section_combo.setStyleSheet("""
            QComboBox {
                font-size: 13px;
                font-weight: bold;
            }
        """)
        button_layout.addWidget(self.section_combo)

        self.show_all_sections_checkbox = QCheckBox("Show all lists")
        self.show_all_sections_checkbox.setChecked(False)
        self.show_all_sections_checkbox.stateChanged.connect(self.update_files_table)
        button_layout.addWidget(self.show_all_sections_checkbox)

        self.add_list_button = QPushButton("Add new list")
        self.add_list_button.clicked.connect(self.add_section)
        button_layout.addWidget(self.add_list_button)

        self.rename_list_button = QPushButton("Rename list")
        self.rename_list_button.clicked.connect(self.rename_list)
        button_layout.addWidget(self.rename_list_button)

        self.remove_list_button = QPushButton("Remove list")
        self.remove_list_button.clicked.connect(self.remove_list)
        button_layout.addWidget(self.remove_list_button)

        self.search_box_saved = QLineEdit()
        self.search_box_saved.setPlaceholderText("Search files...")
        self.search_box_saved.textChanged.connect(self.filter_saved_files)
        button_layout.addWidget(self.search_box_saved)
        
        self.skip_prompt_checkbox = QCheckBox("Don't prompt for source name when saving dragged text")
        self.skip_prompt_checkbox.setChecked(False)  # Optional: set default state
        button_layout.addWidget(self.skip_prompt_checkbox)



        self.move_files_button = QPushButton("Move selected file(s) to new list")
        self.move_files_button.clicked.connect(self.move_files_to_list)
        button_layout.addWidget(self.move_files_button)

        self.remove_saved_file_button = QPushButton("Remove selected file(s)")
        self.remove_saved_file_button.clicked.connect(self.remove_selected_saved_file)
        button_layout.addWidget(self.remove_saved_file_button)

        layout.addLayout(button_layout)

        # Table widget to display all sections and their files
        self.files_table = QTableWidget()
        self.files_table.setColumnCount(6)
        self.files_table.setHorizontalHeaderLabels([
            "Files", "Lists", "Date Modified", "Date Created", "Date Accessed", "Notes"
        ])
        self.files_table.setColumnWidth(0, 580)
        self.files_table.setColumnWidth(1, 100)
        self.files_table.setColumnWidth(2, 120)
        self.files_table.setColumnWidth(3, 120)
        self.files_table.setColumnWidth(4, 120)
        self.files_table.setColumnWidth(5, 600)

        # Enable horizontal scroll for column overflow
        self.files_table.setWordWrap(False)
        self.files_table.setMinimumHeight(0)  # Allow it to shrink fully
        self.files_table.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.files_table.setHorizontalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)

        self.files_table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Interactive)

        self.files_table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
        self.files_table.itemChanged.connect(self.on_note_edited)
        # Vertical splitter for table and preview
        # Create the preview panel
        self.preview_browser = QTextBrowser()
        self.preview_browser.setMinimumHeight(0)  # Allow shrinking
        self.preview_browser.setMaximumHeight(1000)  # Optional limit
        self.preview_browser.setStyleSheet("""
            QTextBrowser {
                font-size: 14px;
                padding: 10px;
                background-color: #f9f9f9;
                border: 1px solid #ccc;
            }
        """)

        # Create vertical splitter to allow resizing
        self.saved_splitter = QSplitter(Qt.Orientation.Vertical)
        self.saved_splitter.addWidget(self.files_table)
        self.saved_splitter.addWidget(self.preview_browser)
        self.saved_splitter.setStretchFactor(0, 1)  # File table
        self.saved_splitter.setStretchFactor(1, 1)  # Preview browser

        # Set initial height ratio: table = 600px, preview = 200px
        self.saved_splitter.setSizes([600, 160])

        # Add splitter to the main layout
        layout.addWidget(self.saved_splitter)
        self.files_table.cellDoubleClicked.connect(self.handle_saved_file_double_click)
        self.files_table.cellClicked.connect(self.preview_saved_file)
        self.files_table.cellDoubleClicked.connect(self.open_saved_file_external)


        # self.files_table.cellDoubleClicked.connect(self.open_saved_file)
        # self.files_table.cellClicked.connect(self.open_saved_file)
        self.files_table.setSortingEnabled(True)
        self.files_table.sortItems(1)
        self.files_table.itemSelectionChanged.connect(self.preview_selected_file)

        self.files_table.setAcceptDrops(True)
        self.files_table.dragEnterEvent = self.dragEnterEvent
        self.files_table.dragMoveEvent = self.dragMoveEvent
        self.files_table.dropEvent = self.dropEvent_saved_files
        self.files_table.mousePressEvent = self.saved_files_mouse_press_event
        self.files_table.mouseMoveEvent = self.saved_files_mouse_move_event
        self._drag_start_position = None
        

        self.tab2.setLayout(layout)

    def preview_selected_file(self):
        selected_items = self.files_table.selectedItems()
        if not selected_items:
            return

        # Get the first column of the selected row
        file_item = next((item for item in selected_items if item.column() == 0), None)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            self.preview_browser.setText("[File does not exist]")
            
            return

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            self.preview_docx_file(file_path)
        elif ext == ".txt":
            self.preview_txt_file(file_path)
        elif ext == ".pdf":
            self.preview_pdf_file(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".gif"]:
            self.preview_image_file(file_path)
    
        else:
            self.preview_browser.setText(f"[Preview not available for file type: {ext}]")
            

    def handle_saved_file_double_click(self, row, column):
        file_item = self.files_table.item(row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' no longer exists.")
            return

        # Optional: show preview for .docx, open others
        if file_path.lower().endswith(".docx"):
            self.preview_docx_file(file_path)
            
        else:
            self.open_file(file_path)
    
    def saved_files_mouse_press_event(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._drag_start_position = event.position().toPoint()
        QTableWidget.mousePressEvent(self.files_table, event)
        
    def preview_saved_file(self, row, column):
        file_item = self.files_table.item(row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            self.preview_browser.setText("[File does not exist]")
            
            return

        ext = os.path.splitext(file_path)[1].lower()

        if ext == ".docx":
            self.preview_docx_file(file_path)
        elif ext == ".txt":
            self.preview_txt_file(file_path)
        elif ext == ".pdf":
            self.preview_pdf_file(file_path)
        elif ext in [".png", ".jpg", ".jpeg", ".gif"]:
            self.preview_image_file(file_path)
    
        else:
            self.preview_browser.setText(f"[Preview not available for file type: {ext}]")
        
        
    def preview_txt_file(self, file_path):
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            self.preview_browser.setText(content or "[File is empty]")
            
        except Exception as e:
            self.preview_browser.setText(f"[Failed to read text file: {e}]")
            
            
    def preview_image_file(self, file_path):
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError("Image file does not exist")

            from PIL import Image

            # Get original image dimensions
            with Image.open(file_path) as img:
                width, height = img.size

            # Calculate scaled dimensions (70%)
            scaled_width = int(width * 0.7)
            scaled_height = int(height * 0.7)

            # Generate HTML with fixed dimensions
            img_html = f'''
                <div align="center">
                    <img src="file:///{file_path}" width="{scaled_width}" height="{scaled_height}" />
                </div>
            '''
            self.preview_browser.setHtml(img_html)
            

        except Exception as e:
            self.preview_browser.setText(f"[Failed to preview image: {e}]")
            

    def preview_pdf_file(self, file_path):
        try:
            import fitz  # PyMuPDF
            doc = fitz.open(file_path)
            content = ""
            for page in doc:
                content += page.get_text()
            self.preview_browser.setText(content.strip() or "[PDF is empty]")
            
        except Exception as e:
            self.preview_browser.setText(f"[Failed to load PDF: {e}]")
            

    def open_saved_file_external(self, row, column):
        file_item = self.files_table.item(row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' no longer exists.")
            return

        self.open_file(file_path)  # Use system viewer
        

    def saved_files_mouse_move_event(self, event):
        if event.buttons() & Qt.MouseButton.LeftButton:
            if self._drag_start_position is None:
                return
            distance = (event.position().toPoint() - self._drag_start_position).manhattanLength()
            if distance < QApplication.startDragDistance():
                return

            selected_rows = set(index.row() for index in self.files_table.selectedIndexes())
            file_paths = []

            for row in selected_rows:
                item = self.files_table.item(row, 0)
                if item:
                    relative_path = item.text().strip()
                    abs_path = os.path.abspath(relative_path)
                    if os.path.exists(abs_path):
                        file_paths.append(abs_path)

            if file_paths:
                urls = [QUrl.fromLocalFile(path) for path in file_paths]

                mime_data = QMimeData()
                mime_data.setUrls(urls)

                drag = QDrag(self.files_table)
                drag.setMimeData(mime_data)

                result = drag.exec(Qt.DropAction.CopyAction)
                print(f"[DEBUG] Drag result: {result}, files dragged: {file_paths}")



    def record_saved_file(self, file_path, section_name):
        """Add an entry to the saved file list for a moved file."""
        with open(self.SAVED_FILES_FILE, "a", encoding="utf-8") as file:
            file.write(f"{file_path}|||{section_name}\n")
            
            
    def dropEvent_saved_files(self, event):
        """Handle drop events in the Saved Files tab."""
        mime = event.mimeData()
        section_name = self.section_combo.currentText()

        if not section_name:
            QMessageBox.warning(self, "No list Selected", "Please select a list to drop files/text into.")
            return

        if mime.hasUrls():
            event.acceptProposedAction()
            move_files = not self.copy_files_checkbox.isChecked()
            for url in mime.urls():
                file_path = url.toLocalFile()
                if os.path.exists(file_path):
                    if move_files:
                        # Implement file move logic if needed
                        pass
                    else:
                        self.add_file_to_section(file_path, section_name)
            self.update_files_table()

        elif (mime.hasText() or 
              mime.hasFormat("text/plain") or 
              mime.hasFormat("text/html") or 
              mime.hasFormat('application/x-qt-windows-mime;value="Text"') or 
              mime.hasFormat('application/x-qt-windows-mime;value="Rich Text Format"')):

            event.acceptProposedAction()
            text = mime.text()

            if mime.hasFormat('application/x-qt-windows-mime;value="Text"'):
                raw_data = mime.data('application/x-qt-windows-mime;value="Text"')
                try:
                    text = str(bytes(raw_data), 'utf-16').strip()
                except Exception:
                    text = raw_data.data().decode(errors="ignore").strip()

            if not text.strip():
                QMessageBox.warning(self, "Empty Drop", "No text found in the dragged content.")
                return

            try:
                import re
                import unicodedata
                from datetime import datetime
                from docx import Document
                from urllib.parse import urlparse

                os.makedirs("lists", exist_ok=True)
                section_folder = os.path.join(self.LISTS_DIR, section_name)
                os.makedirs(section_folder, exist_ok=True)

                # Build timestamp
                now = datetime.now()
                month = now.strftime('%m')
                day = str(now.day)
                #year = str(now.year)
                year = now.strftime('%y')
                hour = str(now.hour % 12 or 12)
                minute = now.strftime('%M')
                ampm = now.strftime('%p')
                timestamp = f"{month}_{day}_{year}_{hour}-{minute}{ampm}"

                # Normalize and extract first 7 words
                normalized_text = unicodedata.normalize("NFC", text)
                first_words = re.findall(r'\w+', normalized_text, re.UNICODE)
                short_title = "_".join(first_words[:7]) if first_words else "untitled"
                short_title = re.sub(r'[\\/*?:"<>|]', '', short_title)
                
                self.debug_clipboard_contents()


                source_hint = self.get_source_from_clipboard_or_prompt(text)


                # ❓ Prompt user if no source found
                if not source_hint:
                    return  # or handle gracefully without asking again


                # Final filename
                filename = f"dragged_text&{source_hint}&{short_title}&{timestamp}.docx"
                file_path = os.path.abspath(os.path.join(section_folder, filename))

                doc = Document()
                doc.add_paragraph(text)
                doc.save(file_path)

                self.record_saved_file(file_path, section_name)
                self.update_files_table()
                print(f"[INFO] dragged text saved to: {file_path}")

            except Exception as e:
                print(f"[ERROR] Saving dragged text failed: {e}")

        else:
            print("[DEBUG] Unknown drag format:", mime.formats())
            event.ignore()





    def get_drag_drop_action(self):
        """Return appropriate Qt drop action based on checkbox state."""
        if hasattr(self, "copy_files_checkbox") and self.copy_files_checkbox.isChecked():
            return Qt.DropAction.CopyAction
        return Qt.DropAction.MoveAction

    def get_file_date(self, file_path, date_type):
        """Get the requested date attribute for a file."""
        try:
            if date_type == "created":
                return os.path.getctime(file_path)  # Creation time
            elif date_type == "modified":
                return os.path.getmtime(file_path)  # Last modified time
            elif date_type == "accessed":
                return os.path.getatime(file_path)  # Last accessed time
        except Exception:
            return 0  # Return 0 if the file doesn't exist or date retrieval fails


    def sort_explorer_files(self):
        """Sort files in the File Explorer by selected date type."""
        if not self.current_directory:
            return  # Don't sort if no directory is loaded

        date_type_mapping = {
            "Sort by Date Created": "created",
            "Sort by Date Modified": "modified",
            "Sort by Date Accessed": "accessed"
        }

        # Get the selected sort type
        selected_sort = self.sort_combo_explorer.currentText()
        date_type = date_type_mapping.get(selected_sort, "modified")  # Default to 'modified' if not found

        # Sort files based on the selected date type
        self.all_files.sort(key=lambda f: self.get_file_date(f, date_type), reverse=True)

        # Refresh the file grid layout with sorted files
        self.refresh_file_grid()


    def refresh_file_grid(self):
        """Rebuilds the file grid after sorting."""
        try:
            self.clear_file_grid()  # Remove all current icons
            row, col = 0, 0

            for file_path in self.all_files:
                if os.path.isdir(file_path):
                    continue
                self.add_file_to_grid(file_path, row, col)
                col += 1
                if col >= 8:
                    col = 0
                    row += 1
        except MemoryError:
            QMessageBox.critical(self, "Memory Error", "Not enough memory to load all files into the grid.")




    def filter_saved_files(self):
        """Filter files in the Saved Files tab based on search query."""
        query = self.search_box_saved.text().lower()
        for row in range(self.files_table.rowCount()):
            file_name = self.files_table.item(row, 0).text().lower()
            self.files_table.setRowHidden(row, query not in file_name)


    def open_saved_file(self, row, column):
        if column != 0:
            return  # Only act when the first column is double-clicked

        file_item = self.files_table.item(row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "Missing File", f"The file '{file_path}' no longer exists.")
            return

        if file_path.lower().endswith(".docx"):
            self.preview_docx_file(file_path)
        else:
            self.preview_browser.setText(f"[Preview not available for file type: {os.path.splitext(file_path)[1]}]")


    def preview_docx_file(self, file_path):
        """Load and display the text content of a .docx file in the preview pane."""
        try:
            from docx import Document
            doc = Document(file_path)
            content = "\n\n".join(paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip())
            self.preview_browser.setText(content or "[Document is empty]")
        except Exception as e:
            self.preview_browser.setText(f"[Failed to preview document: {e}]")


    def on_note_edited(self, item):
        """Handle note editing in the table."""
        if item.column() == 5:  # Only process the Note column
            file_path = os.path.normpath(self.files_table.item(item.row(), 0).text().strip())
            section_name = self.files_table.item(item.row(), 1).text().strip()
            self.notes[(file_path, section_name)] = item.text().strip()
            self.save_notes()






    def add_section(self):
        section_name = self.section_combo.currentText().strip()
        if section_name:
            if section_name not in [self.section_combo.itemText(i) for i in range(self.section_combo.count())]:
                self.section_combo.addItem(section_name)
                self.section_combo_file_explorer.addItem(section_name)  # Update both combo boxes

                self.save_sections()  # Persist changes

                # Set new section as selected and refresh UI
                self.section_combo.setCurrentText(section_name)
                self.section_combo_file_explorer.setCurrentText(section_name)
                self.update_files_table()  
            else:
                QMessageBox.warning(self, "Duplicate list", f"The list '{section_name}' already exists.")
        else:
            QMessageBox.warning(self, "Empty list", "Please enter a valid list name.")


    def rename_list(self):
        current_section = self.section_combo.currentText()
        if not current_section:
            QMessageBox.warning(self, "No list Selected", "Please select a list to rename.")
            return

        new_section_name, ok = QInputDialog.getText(self, "Rename list", "Enter new list name:", text=current_section)
        if ok and new_section_name:
            if new_section_name in [self.section_combo.itemText(i) for i in range(self.section_combo.count())]:
                QMessageBox.warning(self, "Duplicate list", f"The section '{new_section_name}' already exists.")
                return

            # Rename the section in the combo box
            index = self.section_combo.findText(current_section)
            self.section_combo.setItemText(index, new_section_name)


            self.save_sections()  # Save the updated list of sections
            self.update_files_table()  # Refresh the table widget

    def remove_list(self):
        current_section = self.section_combo.currentText()
        if not current_section:
            QMessageBox.warning(self, "No Section Selected", "Please select a section to remove.")
            return

        confirm = QMessageBox.question(self, "Remove Section", f"Are you sure you want to remove the section '{current_section}' and all its files?",
                                     QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
        if confirm == QMessageBox.StandardButton.Yes:
            # Remove the section from the combo box
            index = self.section_combo.findText(current_section)
            self.section_combo.removeItem(index)

            # Delete the section file
            if os.path.exists(f"saved_files_{current_section}.txt"):
                os.remove(f"saved_files_{current_section}.txt")

            self.save_sections()  # Save the updated list of sections
            self.update_files_table()  # Refresh the table widget

    def load_sections(self, load_notes=True):
        """Load sections from the sections.txt file into both combo boxes."""
        self.section_combo.clear()
        self.section_combo_file_explorer.clear()

        if os.path.exists("sections.txt"):
            with open("sections.txt", "r", encoding="utf-8") as file:
                for line in file:
                    section_name = line.strip()
                    if section_name:
                        self.section_combo.addItem(section_name)
                        self.section_combo_file_explorer.addItem(section_name)

        if load_notes:
            self.load_notes()



    def save_sections(self):
        with open("sections.txt", "w", encoding="utf-8") as file:
            for i in range(self.section_combo.count()):
                file.write(self.section_combo.itemText(i) + "\n")

    def update_files_table(self):
        self.files_table.blockSignals(True)
        self.files_table.setSortingEnabled(False)  # 🔻 Disable sorting before changes
        self.files_table.setRowCount(0)

        if self.show_all_sections_checkbox.isChecked():
            for i in range(self.section_combo.count()):
                section_name = self.section_combo.itemText(i)
                self._add_files_from_section(section_name)
        else:
            section_name = self.section_combo.currentText()
            if section_name:
                self._add_files_from_section(section_name)

        self.files_table.setSortingEnabled(True)  # 🔺 Re-enable sorting after populating

        # Enable context menu for copying
        self.files_table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.files_table.customContextMenuRequested.connect(self.show_table_context_menu)

        self.files_table.blockSignals(False)


    def _add_files_from_section(self, section_name):
        if os.path.exists(self.SAVED_FILES_FILE):
            with open(self.SAVED_FILES_FILE, "r", encoding="utf-8") as file:
                for line in file:
                    if "|||" in line:
                        file_path, section = line.strip().split("|||", 1)
                        if section == section_name:
                            row_position = self.files_table.rowCount()
                            self.files_table.insertRow(row_position)
                            full_path = os.path.abspath(file_path.strip())

                            # Column 0 - File Path (Fixed width, tooltip, scrollable)
                            file_item = QTableWidgetItem(full_path)
                            file_item.setToolTip(full_path)  # Show full path on hover
                            file_item.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
                            # file_item.setFlags(file_item.flags() ^ Qt.ItemFlag.ItemIsEditable)  # Disabled edit restriction for File Path
                            self.files_table.setItem(row_position, 0, file_item)

                            # Column 1 - Section name
                            section_item = QTableWidgetItem(section_name)
                            section_item.setFlags(section_item.flags() ^ Qt.ItemFlag.ItemIsEditable)
                            self.files_table.setItem(row_position, 1, section_item)

                            # Column 5 - Notes
                            normalized_path = os.path.normpath(full_path)
                            note_text = self.notes.get((normalized_path, section_name), "")
                            note_item = QTableWidgetItem(note_text)
                            note_item.setFlags(note_item.flags() | Qt.ItemFlag.ItemIsEditable)
                            self.files_table.setItem(row_position, 5, note_item)

                            # Columns 2–4 - Timestamps
                            try:
                                mtime = os.path.getmtime(file_path)
                                ctime = os.path.getctime(file_path)
                                atime = os.path.getatime(file_path)

                                self.files_table.setItem(row_position, 2, QTableWidgetItem(
                                    time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(mtime))
                                ))
                                self.files_table.setItem(row_position, 3, QTableWidgetItem(
                                    time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(ctime))
                                ))
                                self.files_table.setItem(row_position, 4, QTableWidgetItem(
                                    time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(atime))
                                ))
                            except Exception as e:
                                print(f"[ERROR] Could not read timestamps: {file_path} ({e})")





    def load_files_for_section(self, section_name):
        self.update_files_table()

    def move_files_to_list(self):
        selected_items = self.files_table.selectedItems()
        if not selected_items:
            QMessageBox.warning(self, "No Selection", "Please select files to move.")
            return

        # Get the target section
        target_section, ok = QInputDialog.getItem(self, "Move Files", "Select target section:", 
                                                 [self.section_combo.itemText(i) for i in range(self.section_combo.count())], 0, False)
        if not ok or not target_section:
            return

        # Move selected files to the target section
        for item in selected_items:
            if item.column() == 0:  # Only process the file column
                file_path = item.text()
                source_section = self.files_table.item(item.row(), 1).text()

                # Remove the file from the source section
                self.remove_file_from_section(file_path, source_section)

                # Add the file to the target section
                self.add_file_to_section(file_path, target_section)

        self.update_files_table()  # Refresh the table widget

    def remove_file_from_section(self, file_path, section_name):
        if os.path.exists(self.SAVED_FILES_FILE):
            with open(self.SAVED_FILES_FILE, "r", encoding="utf-8") as file:
                lines = file.readlines()

            # Normalize for better matching
            normalized_target = os.path.normcase(os.path.abspath(file_path))

            with open(self.SAVED_FILES_FILE, "w", encoding="utf-8") as file:
                for line in lines:
                    if "|||" not in line:
                        continue
                    saved_path, saved_section = line.strip().split("|||", 1)
                    normalized_saved = os.path.normcase(os.path.abspath(saved_path))

                    if not (normalized_saved == normalized_target and saved_section == section_name):
                        file.write(line)



    

    def add_file_to_section(self, file_path, section_name):
        """Copy file into a section folder and record the entry."""
        if not os.path.exists(file_path) or os.path.isdir(file_path):
            print(f"[SKIP] Not a regular file: {file_path}")
            return


        already_saved = False
        if os.path.exists(self.SAVED_FILES_FILE):
            with open(self.SAVED_FILES_FILE, "r", encoding="utf-8") as file:
                for line in file:
                    saved_path, saved_section = line.strip().split("|||", 1)
                    # Use normalized absolute path for comparison
                    if saved_section == section_name and os.path.basename(saved_path) == os.path.basename(file_path):
                        already_saved = True
                        break

        if already_saved:
            print(f"[INFO] Skipping duplicate: '{file_path}' already saved in list '{section_name}'")
            return

        # Make section folder if it doesn't exist
        section_folder = os.path.join(self.LISTS_DIR, section_name)
        os.makedirs(section_folder, exist_ok=True)

        # Destination path (avoid name collision)
        base_name = os.path.basename(file_path)
        dest_path = os.path.join(section_folder, base_name)
        if os.path.exists(dest_path):
            name, ext = os.path.splitext(base_name)
            dest_path = os.path.join(section_folder, f"{name}_{int(time.time())}{ext}")

        try:
            shutil.copy2(file_path, dest_path)  # Preserve metadata
        except Exception as e:
            print(f"[ERROR] Copying file failed: {e}")
            return

        # Record the original source path (not the destination path)
        with open(self.SAVED_FILES_FILE, "a", encoding="utf-8") as file:
            file.write(f"{dest_path}|||{section_name}\n")




    def load_notes(self):
        self.notes = {}
        if os.path.exists(self.NOTES_FILE):
            with open(self.NOTES_FILE, "r", encoding="utf-8") as file:
                for line in file:
                    try:
                        file_path, section, note = line.split("|||", 2)
                        file_path = os.path.normpath(file_path.strip())
                        section = section.strip()
                        self.notes[(file_path, section)] = note.strip()

                    except ValueError:
                        print(f"[Invalid note line]: {line}")


    def save_notes(self):
        with open(self.NOTES_FILE, "w", encoding="utf-8") as file:
            for (file_path, section), note in self.notes.items():
                file.write(f"{file_path}|||{section}|||{note}\n")



    def open_directory_dialog(self):
        directory = QFileDialog.getExistingDirectory(self, "Select Directory")
        if directory:
            self.current_directory = directory
            self.update_breadcrumb(directory)  # Update breadcrumb path
            self.populate_tree(directory)  
            self.display_files()

            # Update all_files list
            self.all_files = [os.path.join(directory, f) for f in os.listdir(directory)]

            self.display_files()  # Refresh display

    def update_breadcrumb(self, path):
        """Update the breadcrumb navigation with clickable sections."""

        breadcrumb_list = []
        drive, tail = os.path.splitdrive(os.path.abspath(path))
        # Handle Unix root or Windows drive letter as the first breadcrumb
        if drive:
            # For Windows, add drive (e.g., "C:\") as first breadcrumb
            breadcrumb_list.append((drive + os.path.sep, os.path.abspath(drive + os.path.sep)))
        else:
            # For Unix-like systems, if path starts with '/', add '/' as root
            if path.startswith(os.path.sep):
                breadcrumb_list.append((os.path.sep, os.path.abspath(os.path.sep)))

        # Now accumulate the remaining parts
        # Remove any leading os.sep in tail
        tail = tail.lstrip(os.path.sep)
        parts = tail.split(os.path.sep) if tail else []

        accumulated_path = os.path.abspath(drive + os.path.sep) if drive else os.path.abspath(os.path.sep)
        for part in parts:
            accumulated_path = os.path.join(accumulated_path, part)
            breadcrumb_list.append((part, os.path.abspath(accumulated_path)))

        self.path_label.setBreadcrumbs(breadcrumb_list)


    def on_breadcrumb_clicked(self, clicked_path):
        """Handle click event when a breadcrumb section is clicked."""
        from urllib.parse import unquote

        print(f"[ERROR] Raw breadcrumb click href: {clicked_path}")

        clicked_path = unquote(clicked_path)
        if clicked_path.startswith("file:///"):
            clicked_path = clicked_path.replace("file:///", "", 1)
        clicked_path = os.path.normpath(clicked_path)
        print(f"[ERROR] Decoded and normalized path: {clicked_path}")

        if os.path.isdir(clicked_path):
            print(f"[ERROR] Path is a directory. Updating views...")
            self.current_directory = clicked_path
            self.update_breadcrumb(clicked_path)

            self.tree_widget.clear()
            root_item = QTreeWidgetItem([os.path.basename(clicked_path)])
            root_item.setData(0, Qt.ItemDataRole.UserRole, clicked_path)
            # Icon loading deferred for performance
            self.tree_widget.addTopLevelItem(root_item)
            self.populate_subitems(root_item, clicked_path)
            self.tree_widget.expandItem(root_item)

            self.load_thread = FileLoaderThread(clicked_path)
            self.load_thread.files_loaded.connect(self.on_files_loaded)
            self.load_thread.start()

        elif os.path.isfile(clicked_path):
            print(f"[ERROR] Path is a file. Opening file...")
            self.open_file(clicked_path)
        else:
            print(f"[ERROR] Path does not exist or is not accessible: {clicked_path}")



    def populate_tree(self, root_path):
        self.tree_widget.clear()
        root_item = QTreeWidgetItem(self.tree_widget, [os.path.basename(root_path)])
        root_item.setData(0, Qt.ItemDataRole.UserRole, root_path)
        root_item.setIcon(0, self.file_model.fileIcon(self.file_model.index(root_path)))
        self.tree_widget.addTopLevelItem(root_item)
        self.populate_subitems(root_item, root_path)
        self.tree_widget.expandToDepth(0)

    def populate_subitems(self, parent_item, path):
        try:
            for entry in os.scandir(path):
                if entry.is_dir():
                    item = QTreeWidgetItem([entry.name])
                    item.setData(0, Qt.ItemDataRole.UserRole, entry.path)
                    item.setIcon(0, self.file_model.fileIcon(self.file_model.index(entry.path)))

                    # Add dummy child to indicate expandable folder
                    dummy = QTreeWidgetItem(["Loading..."])
                    item.addChild(dummy)

                    parent_item.addChild(item)
        except PermissionError:
            pass



    def on_item_clicked(self, item):
        path = item.data(0, Qt.ItemDataRole.UserRole)

        if not path:
            return

        if os.path.isdir(path):
            # ✅ New behavior: set current directory and update UI
            self.current_directory = path
            self.update_breadcrumb(path)
            self.load_thread = FileLoaderThread(path)
            self.load_thread.files_loaded.connect(self.on_files_loaded)
            self.load_thread.start()
        elif os.path.isfile(path):
            self.open_file(path)




    def open_folder_in_explorer(self, folder_path):
        """Open a folder in Windows File Explorer, macOS Finder, or Linux File Manager."""
        try:
            if sys.platform.startswith("win"):  # Windows 10/11
                print(f"Opening folder in Windows Explorer: {folder_path}")  # ERROR print

                # Best way to open a folder in Windows
                os.startfile(folder_path)

            elif sys.platform.startswith("darwin"):  # macOS
                subprocess.Popen(["open", folder_path])

            elif sys.platform.startswith("linux"):  # Linux
                subprocess.Popen(["xdg-open", folder_path])

            else:
                QMessageBox.critical(self, "Error", "Unsupported OS: Cannot open folder.")

        except Exception as e:
            print(f"Error opening folder: {e}")  # ERRORging print
            QMessageBox.critical(self, "Error", f"Failed to open folder:\n{e}")



    def on_files_loaded(self, files):
        """Handle completion of background file loading."""
        print(f"[ERROR] Background loaded {len(files)} files.")
        self.all_files_full = files
        self.loaded_file_count = 500
        self.all_files = files[:self.loaded_file_count]
        self.display_files()
        

    def display_files(self):
        """Display files based on the selected view mode."""
        self.clear_file_grid()

        if not self.current_directory:
            return  # No directory selected, return early

        # Ensure sorting before displaying files
        selected_sort = self.sort_combo_explorer.currentText()
        date_type_mapping = {
            "Sort by Date Created": "created",
            "Sort by Date Modified": "modified",
            "Sort by Date Accessed": "accessed"
        }
        date_type = date_type_mapping.get(selected_sort, "modified")

        # Sort files based on date
        self.all_files.sort(
            key=lambda f: self.get_file_date(f, date_type) if os.path.exists(f) else 0, 
            reverse=True
        )

        # ERROR: Check if files are present
        #print(f"Displaying {len(self.all_files)} files in {self.current_directory}")

        # Get the selected view mode
        view_mode = self.view_mode_combo.currentText()

        if view_mode == "Icon View":
            self.display_icon_view()
        elif view_mode == "List View":
            self.display_list_view()
        elif view_mode == "Detailed View":
            self.display_detailed_view()



    def display_icon_view(self):
        """Display files in a grid with icons."""
        row, col = 0, 0
        for file_path in self.all_files:
            if os.path.isdir(file_path):
                continue
            self.add_file_to_grid(file_path, row, col)
            col += 1
            if col >= 8:  # Limit columns to 8 before wrapping to a new row
                col = 0
                row += 1


    def display_list_view(self):
        for file_path in self.all_files:
            if os.path.isdir(file_path): continue
            file_frame = QFrame()
            file_frame.setStyleSheet("border: 2px solid transparent;")
            file_layout = QHBoxLayout()

            # File icon
            file_icon = self.file_model.fileIcon(self.file_model.index(file_path))
            icon_label = QLabel()
            icon_label.setPixmap(file_icon.pixmap(24, 24))

            # File name label (clickable)
            file_name = os.path.basename(file_path)
            file_name_label = QLabel(f"<a href='{file_path}'>{file_name}</a>")
            file_name_label.setOpenExternalLinks(False)
            file_name_label.setStyleSheet("color: blue; text-decoration: underline; cursor: pointer;")
            file_name_label.setAlignment(Qt.AlignmentFlag.AlignLeft)
            file_name_label.mousePressEvent = lambda event, path=file_path: self.open_file(path)

            # ✅ Make frame toggle selection on click
            file_frame.mousePressEvent = lambda event, path=file_path, frame=file_frame: self.toggle_file_selection(path, frame)

            # ✅ Apply highlight if already selected
            if file_path in self.selected_files:
                file_frame.setStyleSheet("border: 2px solid blue;")

            file_layout.addWidget(icon_label)
            file_layout.addWidget(file_name_label)
            file_frame.setLayout(file_layout)

            self.grid_layout.addWidget(file_frame)



    def display_detailed_view(self):
        """Display files in a table with details, allowing selection and opening."""
        try:
            table = QTableWidget()
            table.setColumnCount(5)
            table.setHorizontalHeaderLabels([
                "File Name", "Size", "Last Modified", "Date Created", "Date Accessed"
            ])
            table.setColumnWidth(0, 340)
            table.setColumnWidth(1, 75)
            table.setColumnWidth(2, 120)
            table.setColumnWidth(3, 120)
            table.setColumnWidth(4, 120)

            table.mousePressEvent = self.file_table_mouse_press_event
            table.mouseMoveEvent = self.file_table_mouse_move_event
            self._drag_start_position = None

            for file_path in self.all_files:
                if os.path.isdir(file_path):
                    continue
                try:
                    file_name = os.path.basename(file_path)
                    file_size = os.path.getsize(file_path)
                    last_modified = os.path.getmtime(file_path)
                    date_created = os.path.getctime(file_path)
                    date_accessed = os.path.getatime(file_path)

                    row_position = table.rowCount()
                    table.insertRow(row_position)

                    # File name
                    table.setItem(row_position, 0, QTableWidgetItem(file_name))

                    # File size (KB or MB)
                    if file_size < 1024 * 1024:
                        size_display = f"{round(file_size / 1024, 2)} KB"
                    else:
                        size_display = f"{round(file_size / (1024 * 1024), 2)} MB"
                    table.setItem(row_position, 1, QTableWidgetItem(size_display))

                    # File timestamps
                    table.setItem(row_position, 2, QTableWidgetItem(
                        time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(last_modified))))
                    table.setItem(row_position, 3, QTableWidgetItem(
                        time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(date_created))))
                    table.setItem(row_position, 4, QTableWidgetItem(
                        time.strftime('%m-%d-%Y %H:%M:%S', time.localtime(date_accessed))))

                    # Store file path
                    table.item(row_position, 0).setData(Qt.ItemDataRole.UserRole, file_path)

                except (OSError, IOError) as e:
                    print(f"[Warning] Skipping unreadable file: {file_path} ({e})")

            table.itemSelectionChanged.connect(
                lambda: self.update_selected_files_from_table(table))
            table.setSelectionBehavior(QTableWidget.SelectionBehavior.SelectRows)
            table.cellDoubleClicked.connect(self.on_table_file_double_click)
            table.setSortingEnabled(True)
            self.grid_layout.addWidget(table)

        except MemoryError:
            QMessageBox.critical(
                self, "Memory Error", "Not enough memory to render detailed file view.")





    def on_table_file_double_click(self, row, column):
        """Open a file when double-clicked in the Detailed View."""
        table = self.grid_layout.itemAt(0).widget()  # Get the table widget
        file_item = table.item(row, 0)  # Get the first column (file name)

        if file_item:
            file_path = file_item.data(Qt.ItemDataRole.UserRole)  # Retrieve stored file path
            self.open_file(file_path)

    def file_table_mouse_press_event(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self._drag_start_position = event.position().toPoint()
        QTableWidget.mousePressEvent(self.grid_layout.itemAt(0).widget(), event)

    def file_table_mouse_move_event(self, event):
        if event.buttons() & Qt.MouseButton.LeftButton:
            if self._drag_start_position is None:
                return
            distance = (event.position().toPoint() - self._drag_start_position).manhattanLength()
            if distance < QApplication.startDragDistance():
                return

            table = self.grid_layout.itemAt(0).widget()
            selected_rows = list(set(index.row() for index in table.selectedIndexes()))
            file_paths = []

            for row in selected_rows:
                item = table.item(row, 0)
                if item:
                    file_path = item.data(Qt.ItemDataRole.UserRole)
                    if file_path and os.path.exists(file_path):
                        file_paths.append(file_path)

            if file_paths:
                mime_data = QMimeData()
                mime_data.setUrls([QUrl.fromLocalFile(p) for p in file_paths])
                drag = QDrag(table)
                drag.setMimeData(mime_data)
                drag.exec(self.get_drag_drop_action())
                 # ✅ After drag completes, remove files that were moved
                if self.get_drag_drop_action() == Qt.DropAction.MoveAction:
                    self.all_files = [f for f in self.all_files if os.path.exists(f)]
                    self.display_files()


    def add_file_to_grid(self, file_path, row, col):
        file_frame = QFrame()
        file_frame.setFixedSize(100, 100)
        file_frame.setStyleSheet("border: 2px solid transparent;")

        # File icon
        file_icon = self.file_model.fileIcon(self.file_model.index(file_path))
        icon_label = QLabel()
        icon_label.setPixmap(file_icon.pixmap(24, 24))
        icon_label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        # File name label
        file_name = os.path.basename(file_path)
        file_name_label = QLabel(f"<a href='{file_path}'>{file_name}</a>")
        file_name_label.setOpenExternalLinks(False)
        file_name_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        file_name_label.setStyleSheet("color: blue; text-decoration: underline;")
        file_name_label.setCursor(Qt.CursorShape.PointingHandCursor)
        file_name_label.setWordWrap(True)
        file_name_label.linkActivated.connect(lambda: self.open_file(file_path))

        # Selection highlight
        if file_path in self.selected_files:
            file_frame.setStyleSheet("border: 2px solid blue;")

        # Mouse events
        def mouse_press(event):
            file_frame._drag_start_pos = event.position().toPoint()
            self.toggle_file_selection(file_path, file_frame)

        def mouse_move(event):
            if not hasattr(file_frame, "_drag_start_pos"):
                return
            if (event.position().toPoint() - file_frame._drag_start_pos).manhattanLength() < 10:
                return
            mime_data = QMimeData()
            mime_data.setUrls([QUrl.fromLocalFile(file_path)])
            drag = QDrag(file_frame)
            drag.setMimeData(mime_data)
            drag.exec(self.get_drag_drop_action())
            # ✅ After drag completes, remove files that were moved
            if self.get_drag_drop_action() == Qt.DropAction.MoveAction:
                self.all_files = [f for f in self.all_files if os.path.exists(f)]
                self.display_files()


        def mouse_double_click(event):
            if event.button() == Qt.MouseButton.LeftButton:
                self.open_file(file_path)

        file_frame.mousePressEvent = mouse_press
        file_frame.mouseMoveEvent = mouse_move
        file_frame.mouseDoubleClickEvent = mouse_double_click

        # Assemble layout
        layout = QVBoxLayout()
        layout.addWidget(icon_label)
        layout.addWidget(file_name_label)
        file_frame.setLayout(layout)

        # Add to grid
        self.grid_layout.addWidget(file_frame, row, col)



    def open_file(self, file_path):
        try:
            if sys.platform == "win32":
                os.startfile(file_path)  # Windows
            elif sys.platform == "darwin":
                subprocess.Popen(["open", file_path])  # macOS
            else:
                subprocess.Popen(["xdg-open", file_path])  # Linux
        except MemoryError:
            QMessageBox.critical(self, "Memory Error", "Not enough memory to open the file.")
        except Exception as e:
            QMessageBox.critical(self, "Error", f"Failed to open file:\n{e}")


    def toggle_file_selection(self, file_path, frame):
        if file_path in self.selected_files:
            # Remove highlight and remove from selection
            self.selected_files.remove(file_path)
            frame.setStyleSheet("border: 2px solid transparent;")
        else:
            # Add highlight and add to selection
            self.selected_files.append(file_path)
            frame.setStyleSheet("border: 2px solid blue;")

    def save_selected_files(self):
        # Synchronize selection if in detailed view
        view_mode = self.view_mode_combo.currentText()
        if view_mode == "Detailed View":
            table = self.grid_layout.itemAt(0).widget()
            if isinstance(table, QTableWidget):
                self.update_selected_files_from_table(table)

        section_name = self.section_combo_file_explorer.currentText()
        if section_name:
            for file_path in self.selected_files:
                already_saved = False
                if os.path.exists(self.SAVED_FILES_FILE):
                    with open(self.SAVED_FILES_FILE, "r", encoding="utf-8") as f:
                        for line in f:
                            saved_path, saved_section = line.strip().split("|||", 1)
                            if saved_section == section_name and os.path.basename(saved_path) == os.path.basename(file_path):
                                already_saved = True
                                break
                if not already_saved:
                    self.add_file_to_section(file_path, section_name)

            self.selected_files.clear()
            self.clear_file_highlights()
            self.update_files_table()
        else:
            QMessageBox.warning(self, "No Section Selected", "Please select or create a section to save files.")

    

    def update_selected_files_from_table(self, table):
        """Update selected_files list from selected rows in detailed view."""
        self.selected_files.clear()
        selected_ranges = table.selectedRanges()
        for selection in selected_ranges:
            for row in range(selection.topRow(), selection.bottomRow() + 1):
                file_item = table.item(row, 0)
                if file_item:
                    file_path = file_item.data(Qt.ItemDataRole.UserRole)
                    if file_path and os.path.exists(file_path):
                        self.selected_files.append(file_path)

    def clear_file_highlights(self):
        """Clear highlights for all selected files in the file grid."""
        for i in reversed(range(self.grid_layout.count())):
            widget = self.grid_layout.itemAt(i).widget()
            if widget and isinstance(widget, QFrame):
                widget.setStyleSheet("border: 2px solid transparent;")

    def remove_selected_saved_file(self):
        selected_items = self.files_table.selectedItems()
        if selected_items:
            for item in selected_items:
                # Inside the for item in selected_items loop...
                if item.column() == 0:  # Only process the file column
                    original_file_path = item.text()
                    section_name = self.files_table.item(item.row(), 1).text()

                    # Construct the actual saved file path in the "lists/<section>/" folder
                    saved_folder = os.path.join("lists", section_name)
                    saved_file_path = os.path.join(saved_folder, os.path.basename(original_file_path))

                    # If the saved file exists, remove it
                    if os.path.exists(saved_file_path):
                        try:
                            os.remove(saved_file_path)
                        except Exception as e:
                            print(f"[ERROR] Failed to delete saved file: {saved_file_path} - {e}")

                    # Remove entry from the text file and memory
                    self.remove_file_from_section(original_file_path, section_name)

                    note_key = (original_file_path, section_name)
                    if note_key in self.notes:
                        del self.notes[note_key]


            self.update_files_table()
            self.save_notes()


    def add_bookmark(self):
        if not self.current_directory:
            return

        current_path = os.path.normcase(os.path.abspath(self.current_directory))

        # Normalize all current bookmarks
        all_bookmarks = [
            os.path.normcase(os.path.abspath(self.bookmark_list.item(row, col).text()))
            for row in range(self.bookmark_list.rowCount())
            for col in range(self.bookmark_list.columnCount())
            if self.bookmark_list.item(row, col)
        ]

        # Avoid duplicates
        if current_path in all_bookmarks:
            return

        # Add new bookmark
        all_bookmarks.append(current_path)
        self.save_bookmarks_from_list(all_bookmarks)
        self.load_bookmarks()


    def save_bookmarks_from_list(self, bookmarks):
        bookmarks = list(dict.fromkeys(bookmarks))  # remove duplicates while preserving order
        try:
            with open(self.BOOKMARKS_FILE, "w", encoding="utf-8") as file:
                for path in bookmarks:
                    file.write(f"{path}\n")
        except Exception as e:
            print(f"[ERROR] Could not save bookmarks: {e}")

    def remove_selected_bookmark(self):
        selected_items = self.bookmark_list.selectedItems()
        if selected_items:
            for item in selected_items:
                row = self.bookmark_list.row(item)
                col = self.bookmark_list.column(item)
                self.bookmark_list.takeItem(row, col)
            self.save_bookmarks()  # Save after removing a bookmark


    def save_bookmarks(self):
        with open(self.BOOKMARKS_FILE, "w", encoding="utf-8") as file:
            seen = set()
            for row in range(self.bookmark_list.rowCount()):
                for col in range(self.bookmark_list.columnCount()):
                    item = self.bookmark_list.item(row, col)
                    if item and item.text() not in seen:
                        seen.add(item.text())
                        file.write(item.text() + "\n")

    def load_bookmarks(self):
        if not hasattr(self, "bookmark_list"):
            return

        self.bookmark_list.clearContents()
        self.bookmark_list.setRowCount(0)
        self.bookmark_list.setColumnCount(1)
        self.bookmark_list.setHorizontalHeaderLabels(["Bookmarks"])
        self.bookmark_list.verticalHeader().setVisible(False)

        self.bookmark_list.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAsNeeded)
        self.bookmark_list.setHorizontalScrollMode(QAbstractItemView.ScrollMode.ScrollPerPixel)
        self.bookmark_list.setWordWrap(False)
        self.bookmark_list.setSizeAdjustPolicy(QAbstractScrollArea.SizeAdjustPolicy.AdjustToContents)

        all_bookmarks = []
        if os.path.exists(self.BOOKMARKS_FILE):
            with open(self.BOOKMARKS_FILE, "r", encoding="utf-8") as file:
                all_bookmarks = [line.strip() for line in file if line.strip()]

        self.bookmark_list.setRowCount(len(all_bookmarks))

        for row, bookmark in enumerate(all_bookmarks):
            item = QTableWidgetItem(bookmark)
            item.setToolTip(bookmark)
            item.setTextAlignment(Qt.AlignmentFlag.AlignLeft | Qt.AlignmentFlag.AlignVCenter)
            item.setFlags(item.flags() ^ Qt.ItemFlag.ItemIsEditable)
            self.bookmark_list.setItem(row, 0, item)

        # Optional: adjust width of column if needed
        self.bookmark_list.setColumnWidth(0, 500)



    def open_bookmarked_folder(self, item):
        path = item.text()

        if not os.path.exists(path):
            reply = QMessageBox.warning(
                self,
                "Missing Folder",
                f"The folder '{path}' no longer exists.\nWould you like to remove this bookmark?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                row = self.bookmark_list.row(item)
                col = self.bookmark_list.column(item)
                self.bookmark_list.takeItem(row, col)
                self.save_bookmarks()
            return

        self.current_directory = path
        self.update_breadcrumb(path)
        self.populate_tree(path)

        self.load_thread = FileLoaderThread(path)
        self.load_thread.files_loaded.connect(self.on_files_loaded)
        self.load_thread.start()



    def clear_file_grid(self):
        """Remove all file icons from the grid before repopulating."""
        while self.grid_layout.count():
            item = self.grid_layout.takeAt(0)
            if item.widget():
                item.widget().deleteLater()  # Properly delete widgets to free memory


    def dragEnterEvent(self, event):
        mime = event.mimeData()
        if (
            mime.hasUrls() or
            mime.hasText() or
            mime.hasFormat("text/plain") or
            mime.hasFormat("text/html") or
            mime.hasFormat('application/x-qt-windows-mime;value="Text"') or
            mime.hasFormat('application/x-qt-windows-mime;value="Rich Text Format"')
        ):
            event.acceptProposedAction()
        else:
            event.ignore()

    def dragMoveEvent(self, event):
        self.dragEnterEvent(event)  # reuse logic


    def dropEvent(self, event):
        """Handle drop events in the File Explorer tab (grid)."""
        if event.mimeData().hasUrls():
            event.setDropAction(Qt.DropAction.CopyAction)
            event.accept()

            for url in event.mimeData().urls():
                src_path = url.toLocalFile()
                if os.path.isfile(src_path):
                    file_name = os.path.basename(src_path)
                    dest_path = os.path.join(self.current_directory, file_name)

                    # Avoid overwriting existing files
                    if os.path.exists(dest_path):
                        name, ext = os.path.splitext(file_name)
                        dest_path = os.path.join(self.current_directory, f"{name}_{int(time.time())}{ext}")

                    try:
                        shutil.copy2(src_path, dest_path)
                        print(f"[INFO] File copied to: {dest_path}")
                        self.all_files.append(dest_path)
                    except Exception as e:
                        print(f"[ERROR] Failed to copy file: {e}")

            self.display_files()  # Refresh grid
        else:
            event.ignore()
            
    def get_source_from_clipboard_or_prompt(self, dragged_text):
        from urllib.parse import urlparse
        import re

        clipboard = QApplication.clipboard()
        mime = clipboard.mimeData()

        skip_prompt = getattr(self, 'skip_prompt_checkbox', None) and self.skip_prompt_checkbox.isChecked()

        # --- 1. Chromium internal source URL ---
        chromium_url_format = 'application/x-qt-windows-mime;value="Chromium internal source URL"'
        if chromium_url_format in mime.formats():
            try:
                raw_data = mime.data(chromium_url_format)
                url = str(raw_data, "utf-8", errors="ignore").strip()
                if url:
                    domain = urlparse(url).netloc.replace('.', '_')
                    if skip_prompt:
                        return domain
                    preview = f"Detected browser source:\n{url}\n\nUse this as the source?"
                    confirm = QMessageBox.question(self, "Confirm Source", preview,
                                                   QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                    if confirm == QMessageBox.StandardButton.Yes:
                        return domain
            except Exception as e:
                print(f"[ERROR] Failed to read Chromium internal URL: {e}")

        # --- 2. Check for custom extension (SOURCE: ...) ---
        clip_text = mime.text().strip()
        if clip_text.startswith("SOURCE:"):
            try:
                lines = clip_text.splitlines()
                match = re.match(r"SOURCE:\s*(.+?)\s*-\s*(https?://\S+)", lines[0])
                if match:
                    title, url = match.groups()
                    source_body = "\n".join(lines[1:]).strip()
                    if dragged_text.strip()[:30].lower() in source_body.lower():
                        domain = urlparse(url).netloc.replace('.', '_')
                        if skip_prompt:
                            return domain
                        preview = f"Detected Source:\n{title} - {url}\n\nUse this as the source?"
                        confirm = QMessageBox.question(self, "Confirm Source", preview,
                                                       QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
                        if confirm == QMessageBox.StandardButton.Yes:
                            return domain
            except Exception as e:
                print(f"[DEBUG] Failed to parse SOURCE: line: {e}")

        # --- 3. Manual fallback or auto-skip ---
        if skip_prompt:
            return "unspecified_source"

        source_hint, ok = QInputDialog.getText(self, "Enter Source", "Source website or file name:")
        if not ok or not source_hint.strip():
            return "unspecified_source"
        else:
            return re.sub(r'[\\/*?:"<>|]', '', source_hint.strip().replace(' ', '_'))






    def debug_clipboard_contents(self):
        clipboard = QApplication.clipboard()
        mime = clipboard.mimeData()

        print("\n=== Clipboard Debug Info ===")
        for fmt in mime.formats():
            print(f"Format: {fmt}")
            try:
                data = mime.data(fmt)
                print(f"Data: {str(data, 'utf-8', errors='ignore')[:300]}")
            except Exception as e:
                print(f"Could not decode format {fmt}: {e}")
        print("============================\n")
        
        
    def show_table_context_menu(self, position):
        menu = QMenu(self)

        copy_action = QAction("Copy Selected", self)
        copy_action.triggered.connect(self.copy_selected_table_cells)
        menu.addAction(copy_action)

        open_action = QAction("Open File", self)
        open_action.triggered.connect(self.open_selected_file_from_menu)
        menu.addAction(open_action)

        rename_action = QAction("Rename File", self)
        rename_action.triggered.connect(self.rename_selected_file)
        menu.addAction(rename_action)

        # Show non-blocking context menu
        global_pos = self.files_table.viewport().mapToGlobal(position)
        menu.popup(global_pos)

        # Auto-close after 3 seconds
        QTimer.singleShot(3000, menu.close)


    def copy_selected_table_cells(self):
        selection = self.files_table.selectedRanges()
        if not selection:
            return

        file_names = []
        for range_ in selection:
            for row in range(range_.topRow(), range_.bottomRow() + 1):
                item = self.files_table.item(row, 0)
                if item:
                    file_names.append(item.text())

        if not file_names:
            return

        QApplication.clipboard().setText("\n".join(file_names), mode=QClipboard.Mode.Clipboard)

        self.show_temporary_popup(f"Copied {len(file_names)} file name(s)")


    def show_temporary_popup(self, message, duration_ms=3000):
        popup = QFrame(self)
        popup.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.ToolTip)
        popup.setStyleSheet("""
            QFrame {
                background-color: #333;
                color: white;
                border-radius: 8px;
                padding: 10px;
            }
            QLabel {
                color: white;
                font-size: 14px;
            }
        """)
        label = QLabel(message, popup)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)

        popup.resize(250, 60)
        popup.move(self.geometry().center() - popup.rect().center())
        popup.show()

        QTimer.singleShot(duration_ms, popup.close)    
        
    def show_temporary_message(self, text, timeout_ms=3000):
        msg_box = QMessageBox(self)
        msg_box.setWindowTitle("Copied")
        msg_box.setText(text)
        msg_box.setStandardButtons(QMessageBox.StandardButton.NoButton)
        msg_box.setModal(False)
        msg_box.show()

        # Force UI to render and allow timers
        QGuiApplication.processEvents()

        timer = QTimer(self)
        timer.setSingleShot(True)
        timer.timeout.connect(msg_box.close)
        timer.start(timeout_ms)

    def rename_selected_file(self):
        selected_ranges = self.files_table.selectedRanges()
        if not selected_ranges:
            QMessageBox.warning(self, "No Selection", "Please select a file to rename.")
            return

        first_row = selected_ranges[0].topRow()
        file_item = self.files_table.item(first_row, 0)
        section_item = self.files_table.item(first_row, 1)
        if not file_item or not section_item:
            return

        old_path = file_item.text().strip()
        section_name = section_item.text().strip()

        # Use resizable QInputDialog
        dialog = QInputDialog(self)
        dialog.setWindowTitle("Rename File")
        dialog.setLabelText("Enter new file name:")
        dialog.setTextValue(os.path.basename(old_path))
        dialog.resize(750, 100)  # Set width to 750px

        if dialog.exec():
            new_name = dialog.textValue().strip()
            if not new_name:
                return

            new_path = os.path.join(os.path.dirname(old_path), new_name)

            if os.path.exists(new_path):
                QMessageBox.warning(self, "File Exists", f"A file named '{new_name}' already exists.")
                return

            try:
                os.rename(old_path, new_path)
                file_item.setText(new_path)
                self.notes[(os.path.normpath(new_path), section_name)] = self.notes.pop((os.path.normpath(old_path), section_name), "")
                self.save_notes()

                # Update saved_files_all.txt
                if os.path.exists(self.SAVED_FILES_FILE):
                    with open(self.SAVED_FILES_FILE, "r", encoding="utf-8") as f:
                        lines = f.readlines()
                    with open(self.SAVED_FILES_FILE, "w", encoding="utf-8") as f:
                        for line in lines:
                            if line.startswith(old_path + "|||"):
                                f.write(f"{new_path}|||{section_name}\n")
                            else:
                                f.write(line)

                QMessageBox.information(self, "Success", f"File renamed to '{new_name}'.")
            except Exception as e:
                QMessageBox.critical(self, "Rename Failed", f"Could not rename file: {e}")



    def open_selected_file_from_table(self):
        selected_ranges = self.files_table.selectedRanges()
        if not selected_ranges:
            QMessageBox.warning(self, "No Selection", "Please select a file to open.")
            return

        first_row = selected_ranges[0].topRow()
        file_item = self.files_table.item(first_row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            QMessageBox.warning(self, "File Not Found", f"The file '{file_path}' does not exist.")
            return

        try:
            if sys.platform.startswith("win"):
                os.startfile(file_path)
            elif sys.platform.startswith("darwin"):
                subprocess.Popen(["open", file_path])
            elif sys.platform.startswith("linux"):
                subprocess.Popen(["xdg-open", file_path])
            else:
                QMessageBox.warning(self, "Unsupported OS", "Cannot open file on this operating system.")
        except Exception as e:
            QMessageBox.critical(self, "Open Failed", f"Could not open file:{e}")


    def open_selected_file_from_menu(self):
        selected_ranges = self.files_table.selectedRanges()
        if not selected_ranges:
            QMessageBox.warning(self, "No Selection", "Please select a file to open.")
            return

        first_row = selected_ranges[0].topRow()
        file_item = self.files_table.item(first_row, 0)
        if not file_item:
            return

        file_path = file_item.text().strip()
        if not os.path.exists(file_path):
            QMessageBox.critical(self, "File Not Found", f"The file does not exist:{file_path}")
            return

        try:
            if sys.platform.startswith("win"):
                os.startfile(file_path)
            elif sys.platform.startswith("darwin"):
                subprocess.Popen(["open", file_path])
            elif sys.platform.startswith("linux"):
                subprocess.Popen(["xdg-open", file_path])
            else:
                QMessageBox.warning(self, "Unsupported OS", "Cannot open files on this operating system.")
        except Exception as e:
            QMessageBox.critical(self, "Open Failed", f"Could not open the file:{e}")
    


if __name__ == "__main__":
    app = QApplication(sys.argv)
    window = FileExplorerApp()
    window.show()
    sys.exit(app.exec())

    
